import openpyxl  #Importerar saker som behövs behövs
import docx
from docx.enum.text import WD_COLOR_INDEX
import datetime

wb = openpyxl.load_workbook('C:\\Users\\S8hughjo\\Desktop\\PY\\betyg.xlsx')  #Öppnar excel filen
sheet = wb.active                               #Väljer vilken flik vi ska öppna i excel filen
sheet2 = wb.get_sheet_by_name('Kunskapskrav')
sheet3 = wb.get_sheet_by_name('kraven')
sheet4 = wb.get_sheet_by_name('Info')

klass = sheet4.cell(row = 3, column = 2).value

ui = input('Vilken elev vill du fixa ett betyg för?: ')   #En User Input

grade = docx.Document()  #Skapar ett nytt dokument som vi döper till grade

lärare = sheet4.cell(1, 2).value   #Tar ut namnet på läraren och sparar det i en variabel

x = datetime.datetime.now()   #Hämtar dagens datum
dt = x.strftime("%x")

section = grade.sections[0]   #Skapar en parargraph i sektionen högst upp på sidan
header = section.header
paragraph = header.paragraphs[0]
paragraph.text = lärare + "             " + str(dt)

grade.add_heading('Kunskapsmatris Programering 1', 0)  #Lägger till en överskrift

info = grade.add_paragraph()  #skapar en Paragraf 
info.add_run("Klass: " + klass + '        ')  
info.add_run("Namn: " + ui)   #Lägger till text i paragrafen vi nyss skapade

betygs_krav_1 = []   
for i in range(2, 10):      #En for loop som går igenom alla värden på en specifik rad och lägger sedan till det i en array
    betygs_krav_1.append(sheet.cell(row = 1, column = i ).value)

elev = []
for i in range(2, 31):     #Samma sak som ovan fast denna lägger till alla från en column istället
    elev.append(sheet.cell(row = i, column = 1).value)
grade.add_heading('Centralt Innehåll')

runn = []
for i in range(0, len(betygs_krav_1)):    #En for loop som gör igenom alla texter som vi sparade tidigare och skriver sedan ut dem en och en i texten
    lol = grade.add_paragraph(style='List Bullet')
    runn.append(lol.add_run(betygs_krav_1[i]))

mal = []
for i in range(2, 12):    #Samma sak som ovan
    mal.append(sheet2.cell(row = 1, column = i).value)

krav = []
for k in range(2, 5):          #sparar alla krav i en array
    for i in range(2, 12):
        krav.append(sheet3.cell(row = i, column = k).value)
    
grade.add_heading('Kunskapskrav')  #Skriver ut en text som en överskrift

Klarade_A = []
Klarade_C = []
Klarade_E = []

for i in range(1,12):       #Lägger in vilket nummer columnen som eleven hade klarat A/c/e på i varsin array
    namnIndex = elev.index(ui) +2
    if sheet2.cell(row = namnIndex, column = i).value.startswith("A"):
        Klarade_A.append(i-1)
    elif sheet2.cell(row = namnIndex, column = i).value.startswith("C"):
        Klarade_C.append(i-1)
    elif sheet2.cell(row = namnIndex, column = i).value.startswith("E"):
        Klarade_E.append(i-1)

tb = grade.add_table(rows = 1, cols = 4)  #Skapar ett table där vi säger att den ska ha 4 columner och 1 rad
hdr_cells = tb.rows[0].cells
hdr_cells[0].text = 'Rubrik: '   #Ger alla columner namn
hdr_cells[1].text = 'E'
hdr_cells[2].text = 'C'
hdr_cells[3].text = 'A'
for i in range(1, len(mal)+1):   #lägger in texten på varje column
    row_cells = tb.add_row().cells
    row_cells[0].text = mal[i-1]
    row_cells[1].text = krav[i-1]
    row_cells[2].text = krav[(i-1)+len(mal)]
    row_cells[3].text = krav[(i-1)+len(mal)*2]

for k in range(1,len(Klarade_A)+1):  #for loop som går igenom alla värden i klarade a
    a = []
    a.append(tb.cell(Klarade_A[k-1],3))  #Lägger till texten i tablett som stämmer överens med klarade a
    paragraphs = a[0].paragraphs

    for paragraph in paragraphs:        #Gör texten både fet och lägger på grön färg
        for run in paragraph.runs:  
            font = run.font
            font.bold = True
            font.highlight_color = 4

for i in range(1, len(Klarade_C)+1):  #Samma sak som ovan, fast för C
    c = []
    c.append(tb.cell(Klarade_C[i-1],2))
    paragraphs = c[0].paragraphs

    for paragraph in paragraphs:
        for run in paragraph.runs:
            font = run.font
            font.bold = True
            font.highlight_color = 4

for i in range(1, len(Klarade_E)+1):  #Samma sak som ovan fast för E
    e = []
    e.append(tb.cell(Klarade_E[i-1],1))
    paragraphs = e[0].paragraphs

    for paragraph in paragraphs:
        for run in paragraph.runs:
            font = run.font
            font.bold = True
            font.highlight_color = 4


if ui in elev:      #Kollar om User Inputen är med i elev arrayen 
    namnIndex = elev.index(ui) + 2    # skaper en variabel som tar elevens plats i arrayen +2 (array startar på 0 och arrayen tar även med en tom plats)
    for i in range(1, 10):
        if sheet.cell(row = namnIndex , column = i).value.startswith("Genom"):   #Om en cell i excel filen inehåller orden Genom
            runn[i-2].bold = True  #Gör texten fet
    grade.save(r'C:\\Users\\S8hughjo\\Desktop\\PY\\' + str(ui) + '\'s' ' Betyg.docx')  #sparar filen
            
else:
    print('elev finns ej')
